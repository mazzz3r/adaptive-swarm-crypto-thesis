import argparse
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUT = ROOT / "annotation.docx"


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)


def set_run_font(run, size=14, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def style_paragraph(paragraph, *, first_line=True, center=False, bold=False, size=14, after=0):
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(after)
    fmt.first_line_indent = Cm(1.25) if first_line else None
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold)


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, bold=True)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.add_run(text)
    style_paragraph(p)
    return p


def add_reference(doc, text):
    p = doc.add_paragraph()
    p.add_run(text)
    style_paragraph(p, first_line=False)
    return p


def add_page_number_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    p._p.append(fld)


def configure_section(section):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.3)


def parse_args():
    parser = argparse.ArgumentParser(description="Build the Russian thesis annotation DOCX.")
    parser.add_argument(
        "--template",
        type=Path,
        required=True,
        help="Path to the official annotation title-page DOCX template.",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=DEFAULT_OUT,
        help="Output DOCX path. Defaults to annotation.docx in the repository root.",
    )
    return parser.parse_args()


def main():
    args = parse_args()
    doc = Document(args.template)

    for section in doc.sections:
        configure_section(section)

    if doc.tables:
        t = doc.tables[0]
        set_cell_text(t.cell(0, 2), "Безопасные протоколы связи для роевой робототехники", bold=False)
        set_cell_text(t.cell(2, 4), "Бутаков Георгий Игоревич", bold=False)

    for paragraph in doc.paragraphs:
        style_paragraph(paragraph, first_line=False, center=True)

    doc.add_section(WD_SECTION.NEW_PAGE)

    add_heading(doc, "СОДЕРЖАНИЕ")
    contents = [
        ("Введение", "3"),
        ("Краткое описание основной части выпускной квалификационной работы", "3"),
        ("Заключение", "5"),
        ("Список использованной литературы", "5"),
    ]
    for title, page in contents:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.first_line_indent = None
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(f"{title} {' .' * (42 - min(len(title), 42))} {page}")
        set_run_font(r)

    doc.add_section(WD_SECTION.NEW_PAGE)

    add_heading(doc, "ВВЕДЕНИЕ")
    intro = [
        "Выпускная квалификационная работа посвящена разработке и оценке безопасных протоколов связи для роевой робототехники. Актуальность темы определяется тем, что рой автономных агентов зависит от постоянного обмена сообщениями между участниками, но такие сообщения передаются по динамическим беспроводным каналам и обрабатываются устройствами с ограниченными вычислительными и энергетическими ресурсами. Для подобных систем недостаточно выбрать один максимально «тяжелый» криптографический режим на весь период миссии: он может снижать пропускную способность и ускорять расход ресурсов. В то же время постоянно облегченный режим может быть недостаточен при росте угрозы.",
        "Цель работы состоит в проектировании и экспериментальной оценке адаптивного облегченного фреймворка защищенной связи для swarm-oriented систем. Для достижения цели в работе решаются следующие задачи: анализируются архитектуры связи роевых систем и подходы lightweight cryptography; формулируются модель роя, модель угроз и требования к защищенному каналу; сравниваются базовые криптографические примитивы; проектируется функция адаптивного выбора профиля; реализуется воспроизводимый pairwise benchmark; проводится сравнение режимов heavy, balanced, lightweight и adaptive при ограничениях CPU, rekeying и сетевых помехах.",
        "Объектом исследования является защищенная коммуникация в распределенных роевых системах. Предметом исследования является адаптивный выбор криптографического профиля для защищенной data plane. Научная новизна работы заключается в объединении формальной модели адаптивной связи с исполняемым прототипом, который позволяет измерить компромисс между уровнем защиты и стоимостью выполнения, а не описывать адаптацию только на концептуальном уровне.",
    ]
    for text in intro:
        add_body(doc, text)

    add_heading(doc, "КРАТКОЕ ОПИСАНИЕ ОСНОВНОЙ ЧАСТИ ВЫПУСКНОЙ КВАЛИФИКАЦИОННОЙ РАБОТЫ")
    main_part = [
        "В первой содержательной части работы выполнен обзор литературы по коммуникационным архитектурам роевых систем, lightweight cryptography, децентрализованному доверию и адаптивным security frameworks. Анализ показывает, что mesh, star и cluster-based топологии по-разному влияют на отказоустойчивость и управление ключами. Работы по Swarm Relays, RASS, blockchain-based coordination и firmware attestation демонстрируют важность устойчивой инфраструктуры, но не закрывают полностью задачу легкой и динамически настраиваемой защиты сообщений. Поэтому в работе выделены три разрыва: статичность криптографических конфигураций, раздельное рассмотрение связности, trust layer и message protection, а также недостаток исполняемых сравнений между статическими и адаптивными профилями.",
        "Далее в работе задается формальная модель роя как динамического графа G(V, E, t), где вершины соответствуют агентам, а ребра описывают доступные каналы связи в конкретный момент времени. Состояние агента включает позицию, скорость, запас энергии, вычислительные ресурсы и текущие криптографические параметры. На этой основе определяется функция адаптации F: S(t) -> P, которая сопоставляет наблюдаемому состоянию системы набор параметров защиты. Модель угроз учитывает пассивного противника, активное изменение и инъекцию сообщений, replay, impersonation, Sybil, wormhole и denial-of-service сценарии. Ключевыми требованиями к data plane являются конфиденциальность, аутентификация, целостность и ограничение последствий компрометации через обновление ключевого материала.",
        "Сравнение базовых протоколов показывает, что один примитив не решает всю задачу. X25519 используется для pairwise session establishment, HKDF-SHA256 - для вывода epoch-specific ключевого материала, AEAD-схемы - для защиты полезной нагрузки. AES-GCM применяется в heavy и balanced профилях, ChaCha20-Poly1305 - в lightweight профиле. Дополнительная runtime-аутентификация в heavy и balanced профилях строится на HMAC-SHA256, а в lightweight профиле используется sequence-bound hash token. Такая гибридная архитектура отделяет дорогой этап установления ключей от более дешевой защиты каждого сообщения.",
        "Практическая составляющая работы построена как последовательная реализация трех дискретных профилей и адаптивного выбора между ними. Heavy profile использует AES-256-GCM и HMAC-SHA256; balanced profile использует AES-192-GCM и HMAC-SHA256; lightweight profile использует ChaCha20-Poly1305 и sequence-bound hash token. Threshold-based функция выбора работает следующим образом: при threat >= 0.8 система выбирает heavy profile, при энергии ниже 0.2Emax выбирает lightweight profile, в остальных случаях работает balanced profile. Алгоритм не заявляется как глобально оптимальный, но достаточен для проверки усиления защиты при росте угрозы и снижения стоимости при низком энергетическом бюджете.",
        "Для проверки подхода реализован контейнерный benchmark с двумя peer-узлами, соединенными через gRPC. Основная серия проводилась при ограничении 0.05 CPU core на узел, лимите памяти 128 MiB, размере сообщения 65 536 bytes, продолжительности 60 seconds и rekey interval 2 seconds. В итоговой серии собрано 228 summary rows, включая main stability sweep, CPU sensitivity, rekey sensitivity и network impairment. Метрики включают latency, throughput, delivered-message count, drop rate, number of switches, profile mix, normalized energy proxy и cryptographic work per delivered message.",
        "Результаты показывают, что при жестком CPU cap система достигает saturation между 25 и 75 messages per second. Наиболее устойчивым default в условиях вычислительного ограничения оказался lightweight mode: он показывает минимальную cryptographic work per delivered message, около 0.125 MiB, тогда как heavy и balanced обрабатывают около 0.250 MiB на доставленное сообщение из-за дополнительного HMAC pass. Adaptive mode располагается между статическими крайностями и в основной серии дает около 0.224 MiB cryptographic work per delivered message.",
        "Отдельно проверено adaptive switching behavior. В репрезентативном clean 100 Hz запуске система зафиксировала четыре ожидаемых перехода: в heavy mode примерно на 12-й секунде, обратно в balanced на 24-й секунде, в lightweight на 36-й секунде и снова в heavy на 48-й секунде. Среди 57 adaptive runs в 54 запусках записаны все ожидаемые переходы; исключения относятся к deliberately severe lossy-network scenario. Следовательно, adaptive mode корректно реализует заданную политику в нормальных условиях, но не должен описываться как универсально лучший по всем метрикам.",
        "Sensitivity analysis уточняет границы результата. При 0.05 core и 200 Hz потери составляют примерно 62-69%, при 0.10 core ситуация заметно улучшается, а при 0.25 core все режимы выдерживают 200 Hz без потерь. Изменение rekey interval не становится доминирующим фактором при основном CPU bottleneck. Network impairment experiments показывают, что часть ограничений связана не с криптографией как таковой, а с transport-level fragility контейнерного gRPC стенда.",
    ]
    for text in main_part:
        add_body(doc, text)

    add_heading(doc, "ЗАКЛЮЧЕНИЕ")
    conclusion = [
        "В результате работы разработан и экспериментально оценен адаптивный фреймворк защищенной связи для роевой робототехники. Основной вклад состоит в том, что формальная модель и архитектура не остались только теоретическими: они были переведены в воспроизводимый pairwise prototype с измеримыми режимами heavy, balanced, lightweight и adaptive. Полученные результаты подтверждают, что облегченный профиль является наиболее эффективным default при доминирующем вычислительном ограничении, а adaptive profile полезен как policy-driven compromise, который направляет более дорогую защиту в периоды повышенной угрозы и снижает стоимость при низком энергетическом бюджете.",
        "Практическая значимость работы заключается в возможности параметризовать один и тот же pattern переключения для различных миссий: disaster response, agricultural robotics, low-bandwidth links и adversarial UAV operation. При этом выводы ограничены областью эксперимента. Benchmark подтверждает поведение constrained host-based pairwise prototype, но не заменяет проверку на физических роботах, в large-swarm topology, при реальном sensing угроз и с полноценной authenticated control plane. Перспективы дальнейшей работы включают аутентификацию bootstrap и profile-update сообщений, multi-peer validation, hardware evaluation и последующую оценку post-quantum вариантов session establishment.",
    ]
    for text in conclusion:
        add_body(doc, text)

    add_heading(doc, "СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ")
    refs = [
        "1. Dias P. G. F., Silva M. C., Rocha Filho G. P., Vargas P. A., Cota L. P., Pessin G. Swarm Robotics: A Perspective on the Latest Reviewed Concepts and Applications // Sensors. 2021. Vol. 21, No. 6. Article 2062.",
        "2. Debie E., Kasmarik K., Garratt M. Swarm Robotics: A Survey from a Multi-Tasking Perspective // ACM Computing Surveys. 2023. Vol. 56, No. 2. P. 1-38.",
        "3. Varadharajan V. S., St-Onge D., Adams B., Beltrame G. Swarm Relays: Distributed Self-Healing Ground-and-Air Connectivity Chains // IEEE Robotics and Automation Letters. 2020. Vol. 5, No. 4. P. 5347-5354.",
        "4. Chen L., Ng S.-L. Securing Emergent Behaviour in Swarm Robotics // Journal of Information Security and Applications. 2022. Vol. 64. Article 103047.",
        "5. Han P., Sui A., Wu J. Lightweight Secure Communication Supporting Batch Authentication for UAV Swarm // Drones. 2025. Vol. 9, No. 2. Article 139.",
        "6. Adeniyi A. E., Jimoh R. G., Awotunde J. B. A Systematic Review on Elliptic Curve Cryptography Algorithm for Internet of Things // Computers and Electrical Engineering. 2024. Vol. 118. Article 109330.",
        "7. Manikandan K., Sriramulu R. ASMTP: Anonymous Secure Messaging Token-Based Protocol Assisted Data Security in Swarm of Unmanned Aerial Vehicles // International Journal of Network Management. 2024. Vol. 34, No. 6. Article e2271.",
        "8. RFC 7748. Elliptic Curves for Security. IRTF, 2016; RFC 8439. ChaCha20 and Poly1305 for IETF Protocols. IETF, 2018.",
    ]
    for ref in refs:
        add_reference(doc, ref)

    for i, section in enumerate(doc.sections):
        configure_section(section)
        section.different_first_page_header_footer = i == 0
        if i > 0:
            section.footer.is_linked_to_previous = False

    # Keep the title page visually unnumbered. Subsequent sections carry page numbers.
    for section in doc.sections[1:]:
        add_page_number_footer(section)

    doc.save(args.output)
    print(args.output)


if __name__ == "__main__":
    main()
